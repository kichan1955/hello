"""
Convert 1percent_plus_book.html → 1percent_plus_book.docx
- Extracts each section's SVG illustration and converts to PNG
- Preserves all original text with heading styles
- Inserts half-page illustrations at the start of each section
"""

import re, io, os, html as unescape_mod
from bs4 import BeautifulSoup
import cairosvg
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── load HTML ──────────────────────────────────────────────────────────────
with open('/home/user/hello/1percent_plus_book.html', encoding='utf-8') as f:
    html = f.read()

soup = BeautifulSoup(html, 'html.parser')

# ── helpers ─────────────────────────────────────────────────────────────────
def svg_to_png_bytes(svg_str: str, width=1400) -> bytes:
    """Render SVG string → PNG bytes at given pixel width."""
    # cairosvg requires explicit width/height when only viewBox is set
    height = int(width * 380 / 800)
    svg_str = re.sub(r'(<svg\b)', f'\\1 width="{width}" height="{height}"', svg_str, count=1)
    return cairosvg.svg2png(bytestring=svg_str.encode('utf-8'), output_width=width)

def set_font(run, name='나눔명조', size_pt=11, bold=False,
             color: RGBColor = None, italic=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Korean fallback
    r = run._r
    rpr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '나눔명조')
    rpr.insert(0, rFonts)

def para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6,
             left_indent=0):
    fmt = para.paragraph_format
    fmt.alignment = align
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if left_indent:
        fmt.left_indent = Cm(left_indent)

NAVY      = RGBColor(0x0d, 0x1b, 0x2a)
GOLD      = RGBColor(0xc9, 0xa8, 0x4c)
GOLD_LIGHT= RGBColor(0xe0, 0xc5, 0x7a)
MUTED     = RGBColor(0x55, 0x57, 0x70)
BODY_TEXT = RGBColor(0x2a, 0x2a, 0x3e)

# ── Document setup ───────────────────────────────────────────────────────────
doc = Document()

# Page size A4, margins
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(3)
section.top_margin  = section.bottom_margin = Cm(2.5)

# Normal style
normal = doc.styles['Normal']
normal.font.name = '나눔명조'
normal.font.size = Pt(11)

# ── cover page ───────────────────────────────────────────────────────────────
def add_cover(doc):
    for _ in range(6): doc.add_paragraph()
    badge = doc.add_paragraph()
    para_fmt(badge, WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    r = badge.add_run('── 대한민국 히든챔피언 자서전 ──')
    set_font(r, size_pt=10, color=GOLD)

    title = doc.add_paragraph()
    para_fmt(title, WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
    r = title.add_run('1% Plus')
    set_font(r, size_pt=52, bold=True, color=GOLD)

    sub = doc.add_paragraph()
    para_fmt(sub, WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    r = sub.add_run('1%를 향한 여정  ·  세상이 모르는 세계 1위의 비밀')
    set_font(r, size_pt=13, color=GOLD_LIGHT)

    div = doc.add_paragraph()
    para_fmt(div, WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    r = div.add_run('─' * 28)
    set_font(r, size_pt=10, color=GOLD)

    desc1 = doc.add_paragraph()
    para_fmt(desc1, WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    r = desc1.add_run('산업용 모니터 세계 1위 코텍 창업자')
    set_font(r, size_pt=12, color=MUTED)

    desc2 = doc.add_paragraph()
    para_fmt(desc2, WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    r = desc2.add_run('이한구 회장')
    set_font(r, size_pt=16, bold=True, color=NAVY)
    r2 = desc2.add_run('의 도전 · 열정 · 신뢰')
    set_font(r2, size_pt=12, color=MUTED)

    for _ in range(4): doc.add_paragraph()

    author = doc.add_paragraph()
    para_fmt(author, WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    r = author.add_run('박기찬 저')
    set_font(r, size_pt=13, bold=True, color=NAVY)

    inst = doc.add_paragraph()
    para_fmt(inst, WD_ALIGN_PARAGRAPH.CENTER)
    r = inst.add_run('인하대학교 경영대학 명예교수  |  aSSIST 석좌교수')
    set_font(r, size_pt=10, color=MUTED)

    doc.add_page_break()

add_cover(doc)

# ── TOC page ─────────────────────────────────────────────────────────────────
toc_title = doc.add_paragraph()
para_fmt(toc_title, WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
r = toc_title.add_run('목  차')
set_font(r, size_pt=24, bold=True, color=NAVY)

toc_sub = doc.add_paragraph()
para_fmt(toc_sub, WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
r = toc_sub.add_run('Table of Contents')
set_font(r, size_pt=10, color=GOLD)

toc_entries = [
    ('추천의 글', ''),
    ('저자서문', '히든 챔피언의 히든 스토리'),
    ('프롤로그', '세상을 바꾼 하나의 질문'),
    ('제 1 장', '가난이 단련시킨 창업정신'),
    ('제 2 장', '동전 하나의 혁명 — 껌 자동판매기'),
    ('제 3 장', '부도의 아픔 속에서 얻은 평생의 철학'),
    ('제 4 장', '브라운관에서 산업용 모니터로'),
    ('제 5 장', '배반과 질병, 두 개의 전선'),
    ('제 6 장', 'Free Voltage와 글로벌 돌파'),
    ('제 7 장', 'IGT와의 첫 접촉'),
    ('제 8 장', '100% 리콜 — 신뢰를 파는 방법'),
    ('제 9 장', '코스닥 상장 — 세상에 얼굴을 내밀다'),
    ('제 10 장', 'Auto Color Bias — 最·知·信 경영철학'),
    ('제 13·14 장', '인사철학 · 조직문화 — 사람이 전략이다'),
    ('제 15 장', '승계의 철학 — 기업은 기업의 것이다'),
    ('제 16 장', '김영달과의 M&A — 완벽한 승계'),
    ('에필로그', 'SER-M으로 보는 이한구 경영학'),
]

for num, title in toc_entries:
    entry = doc.add_paragraph()
    para_fmt(entry, WD_ALIGN_PARAGRAPH.LEFT, space_before=1, space_after=4, left_indent=0.5)
    r1 = entry.add_run(f'{num}  ')
    set_font(r1, size_pt=10, bold=True, color=GOLD)
    r2 = entry.add_run(title if title else '')
    set_font(r2, size_pt=10, color=NAVY)

doc.add_page_break()

# ── section builder ───────────────────────────────────────────────────────────
def add_section_heading(doc, label: str, title: str, subtitle: str = ''):
    if label:
        lbl_p = doc.add_paragraph()
        para_fmt(lbl_p, WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4)
        r = lbl_p.add_run(label)
        set_font(r, size_pt=9, bold=True, color=GOLD)

    h = doc.add_paragraph()
    para_fmt(h, WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=6)
    r = h.add_run(title)
    set_font(r, size_pt=22, bold=True, color=NAVY)

    if subtitle:
        s = doc.add_paragraph()
        para_fmt(s, WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=10)
        r = s.add_run(subtitle)
        set_font(r, size_pt=11, italic=True, color=MUTED)

    # gold rule
    rule = doc.add_paragraph()
    para_fmt(rule, WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=16)
    r = rule.add_run('─' * 42)
    set_font(r, size_pt=8, color=GOLD)


def add_illustration(doc, svg_str: str, caption: str = ''):
    """Render SVG → PNG and insert as half-page image."""
    png = svg_to_png_bytes(svg_str, width=1400)
    img_stream = io.BytesIO(png)
    # Half page width ≈ usable width = A4 21cm - 6cm margins = 15cm
    img_para = doc.add_paragraph()
    para_fmt(img_para, WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
    run = img_para.add_run()
    run.add_picture(img_stream, width=Cm(15))
    if caption:
        cap = doc.add_paragraph()
        para_fmt(cap, WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)
        r = cap.add_run(caption)
        set_font(r, size_pt=8, italic=True, color=MUTED)


def add_body_text(doc, elements):
    """Add body paragraphs from BeautifulSoup elements."""
    for el in elements:
        tag = el.name if hasattr(el, 'name') else None
        text = el.get_text(strip=True) if tag else ''
        if not text:
            continue

        if tag == 'h3' and 'sub-heading' in el.get('class', []):
            p = doc.add_paragraph()
            para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=5, left_indent=0.3)
            r = p.add_run(text)
            set_font(r, size_pt=12, bold=True, color=NAVY)

        elif tag == 'h4' and 'box-heading' in el.get('class', []):
            p = doc.add_paragraph()
            para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
            r = p.add_run(text)
            set_font(r, size_pt=11, bold=True, color=RGBColor(0x2e,0x4a,0x6a))

        elif tag == 'p' and 'practice' in el.get('class', []):
            p = doc.add_paragraph()
            para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=8, left_indent=0.5)
            r = p.add_run(text)
            set_font(r, size_pt=10.5, italic=True, color=GOLD)

        elif tag == 'p' and 'quote-attr' in el.get('class', []):
            p = doc.add_paragraph()
            para_fmt(p, WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=8)
            r = p.add_run(text)
            set_font(r, size_pt=10.5, italic=True, color=MUTED)

        elif tag == 'p' and 'principle-item' in el.get('class', []):
            p = doc.add_paragraph()
            para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=4, left_indent=0.8)
            r = p.add_run(text)
            set_font(r, size_pt=10.5, color=NAVY)

        elif tag == 'p' and 'data-line' in el.get('class', []):
            p = doc.add_paragraph()
            para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=3, left_indent=0.5)
            r = p.add_run(text)
            set_font(r, size_pt=10.5, color=RGBColor(0x2e,0x4a,0x6a))

        elif tag in ('p', None):
            p = doc.add_paragraph()
            para_fmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=7)
            r = p.add_run(text)
            set_font(r, size_pt=11, color=BODY_TEXT)

# ── Parse HTML sections ───────────────────────────────────────────────────────
def get_section_data(section_el):
    """Return (svg_str, body_elements) from a section element."""
    fig = section_el.find('figure', class_='illustration')
    svg_str = str(fig.find('svg')) if fig else ''
    body_div = section_el.find('div', class_='chapter-body')
    body_els = list(body_div.children) if body_div else []
    return svg_str, body_els

# Process each section
sections_config = [
    # (html_id, label, title, subtitle, page_break_after)
    ('recommend',  '',            '추천의 글',                '',                                           True),
    ('preface',    '',            '저자서문',                  '히든 챔피언의 히든 스토리',                    True),
    ('prologue',   '프롤로그',    '세상을 바꾼 하나의 질문',    '',                                           True),
    ('ch1',        '제 1 장',     '가난이 단련시킨 창업정신',   '',                                           True),
    ('ch2',        '제 2 장',     '동전 하나의 혁명',           '껌 자동판매기',                               True),
    ('ch3',        '제 3 장',     '부도의 아픔 속에서 얻은 평생의 철학', '',                                   True),
    ('ch4',        '제 4 장',     '브라운관에서 산업용 모니터로','',                                          True),
    ('ch5',        '제 5 장',     '배반과 질병, 두 개의 전선',  '',                                           True),
    ('ch6',        '제 6 장',     'Free Voltage와 글로벌 돌파', '',                                           True),
    ('ch7',        '제 7 장',     'IGT와의 첫 접촉',            '',                                           True),
    ('ch8',        '제 8 장',     '100% 리콜 — 신뢰를 파는 방법','',                                         True),
    ('ch9',        '제 9 장',     '코스닥 상장 — 세상에 얼굴을 내밀다','',                                    True),
    ('ch10',       '제 10 장',    'Auto Color Bias',            '最·知·信 경영철학 · 1% Plus의 본질',         True),
    ('ch13',       '제 13·14 장', '인사철학 · 조직문화',        '사람이 전략이다',                             True),
    ('ch15',       '제 15 장',    '승계의 철학',                '기업은 기업의 것이다',                         True),
    ('ch16',       '제 16 장',    '김영달과의 M&A — 완벽한 승계','',                                         True),
    ('epilogue',   '에필로그',    'SER-M으로 보는 이한구 경영학','',                                          True),
    ('q100',       '',            '창업가를 위한 핵심 원칙',     '',                                           False),
]

for i, (sec_id, label, title, subtitle, page_brk) in enumerate(sections_config):
    el = soup.find(id=sec_id)
    if not el:
        print(f'  [skip] #{sec_id} not found')
        continue

    print(f'Processing #{sec_id} …')
    svg_str, body_els = get_section_data(el)

    add_section_heading(doc, label, title, subtitle)

    if svg_str:
        add_illustration(doc, svg_str, caption=f'{label}  {title}' if label else title)

    add_body_text(doc, body_els)

    if page_brk:
        doc.add_page_break()

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/home/user/hello/1percent_plus_book.docx'
doc.save(out)
print(f'\n✓ Saved: {out}')
import os
print(f'  Size: {os.path.getsize(out):,} bytes ({os.path.getsize(out)//1024} KB)')
